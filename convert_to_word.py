#!/usr/bin/env python3
"""
Convertidor de reportes técnicos de Markdown a Word con imágenes incrustadas
para proyectos de análisis estructural naval.
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import matplotlib.pyplot as plt
import pandas as pd
from PIL import Image
import io

class NavalReportConverter:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """Configura estilos personalizados para reportes técnicos navales"""
        styles = self.doc.styles
        
        # Estilo para títulos principales
        title_style = styles.add_style('NavalTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Estilo para subtítulos
        heading_style = styles.add_style('NavalHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        
        # Estilo para texto normal
        normal_style = styles.add_style('NavalText', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(11)
        
        # Estilo para tablas técnicas
        table_style = styles.add_style('NavalTable', WD_STYLE_TYPE.PARAGRAPH)
        table_style.font.name = 'Arial'
        table_style.font.size = Pt(9)
        
    def convert_markdown_to_word(self, md_file_path, output_path=None):
        """Convierte archivo markdown a Word con imágenes y tablas"""
        if output_path is None:
            output_path = md_file_path.replace('.md', '.docx')
            
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # Procesar línea por línea
        lines = content.split('\n')
        current_table = []
        in_table = False
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
                
            # Procesar encabezados
            if line.startswith('#'):
                self.process_header(line)
                
            # Procesar tablas
            elif '|' in line and not line.startswith('!'):
                if not in_table:
                    in_table = True
                    current_table = []
                current_table.append(line)
                
            # Procesar imágenes
            elif '![' in line and '](' in line:
                if in_table:
                    self.create_table(current_table)
                    current_table = []
                    in_table = False
                self.process_image(line, os.path.dirname(md_file_path))
                
            # Procesar texto normal
            elif not line.startswith('---') and not line.startswith('**Archivos:') and not line.startswith('- Herramienta:'):
                if in_table:
                    self.create_table(current_table)
                    current_table = []
                    in_table = False
                self.process_text(line)
                
        # Crear tabla final si queda alguna
        if in_table and current_table:
            self.create_table(current_table)
            
        self.doc.save(output_path)
        return output_path
        
    def process_header(self, line):
        """Procesa encabezados markdown"""
        level = len(line.split()[0])
        text = line.lstrip('#').strip()
        
        if level == 1:
            para = self.doc.add_paragraph(text, style='NavalTitle')
        elif level == 2:
            para = self.doc.add_paragraph(text, style='NavalHeading')
        else:
            para = self.doc.add_paragraph(text)
            para.runs[0].bold = True
            
    def process_text(self, line):
        """Procesa texto normal con formato"""
        if '**' in line:
            # Texto con negritas
            parts = line.split('**')
            para = self.doc.add_paragraph()
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Partes impares son negritas
                    run = para.add_run(part)
                    run.bold = True
                else:
                    para.add_run(part)
        else:
            self.doc.add_paragraph(line, style='NavalText')
            
    def process_image(self, line, base_path):
        """Procesa imágenes markdown"""
        # Extraer ruta de imagen
        match = re.search(r'!\[.*?\]\((.*?)\)', line)
        if match:
            image_path = match.group(1)
            full_path = os.path.join(base_path, image_path)
            
            if os.path.exists(full_path):
                try:
                    self.doc.add_picture(full_path, width=Inches(6))
                    # Agregar leyenda
                    caption_match = re.search(r'!\[(.*?)\]', line)
                    if caption_match and caption_match.group(1):
                        caption = self.doc.add_paragraph(
                            f"Figura: {caption_match.group(1)}", 
                            style='NavalText'
                        )
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    self.doc.add_paragraph(f"[Error al cargar imagen: {image_path}]")
                    
    def create_table(self, table_lines):
        """Crea tabla Word desde líneas markdown"""
        if len(table_lines) < 2:  # Necesita al menos encabezado y separador
            return
            
        # Parsear filas
        rows = []
        for i, line in enumerate(table_lines):
            if i == 1:  # Saltar línea separadora
                continue
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                rows.append(cells)
                
        if not rows:
            return
            
        # Crear tabla
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        
        # Llenar tabla
        for i, row_data in enumerate(rows):
            for j, cell_data in enumerate(row_data):
                if i == 0:  # Encabezado
                    cell = table.cell(i, j)
                    cell.text = cell_data
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    table.cell(i, j).text = cell_data
                    
        # Agregar espacio después de tabla
        self.doc.add_paragraph()
        
    def add_technical_summary(self, summary_data):
        """Agrega resumen técnico con datos del proyecto"""
        self.doc.add_paragraph("RESUMEN TÉCNICO", style='NavalHeading')
        
        # Crear tabla de resumen
        table = self.doc.add_table(rows=len(summary_data), cols=2)
        table.style = 'Table Grid'
        
        for i, (key, value) in enumerate(summary_data.items()):
            table.cell(i, 0).text = key
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
            table.cell(i, 1).text = str(value)
            
    def add_compliance_table(self, compliance_data):
        """Agrega tabla de cumplimiento normativo"""
        self.doc.add_paragraph("VERIFICACIÓN DE CUMPLIMIENTO NORMATIVO", style='NavalHeading')
        
        # Crear tabla de cumplimiento
        headers = ["Requisito", "Valor Calculado", "Valor Requerido", "Estado"]
        table = self.doc.add_table(rows=len(compliance_data) + 1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Encabezados
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            
        # Datos
        for i, (req, calc, req_val, status) in enumerate(compliance_data):
            table.cell(i+1, 0).text = req
            table.cell(i+1, 1).text = str(calc)
            table.cell(i+1, 2).text = str(req_val)
            table.cell(i+1, 3).text = status
            
            # Colorear estado
            if "CUMPLE" in status:
                table.cell(i+1, 3).paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
            else:
                table.cell(i+1, 3).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)

def main():
    """Función principal para convertir reportes"""
    converter = NavalReportConverter()
    
    # Convertir reportes principales
    reports = [
        "ENTREGA 4/RESUMEN_EJECUTIVO.md",
        "ENTREGA 4/RESUMEN_TECNICO_FINAL.md", 
        "ENTREGA 4/REPORTE_CUADERNA_MAESTRA.md",
        "ENTREGA 4/DOCUMENTO_ENTREGA_FINAL.md"
    ]
    
    for report_path in reports:
        if os.path.exists(report_path):
            print(f"Convirtiendo: {report_path}")
            try:
                output_path = converter.convert_markdown_to_word(report_path)
                print(f"✓ Convertido a: {output_path}")
            except Exception as e:
                print(f"✗ Error al convertir {report_path}: {e}")
        else:
            print(f"Archivo no encontrado: {report_path}")

if __name__ == "__main__":
    main()